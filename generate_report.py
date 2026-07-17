#!/usr/bin/env python3
"""生成XXE漏洞检测与修复报告"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def add_code(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0xF5F5F5"/>')
    p._p.get_or_add_pPr().append(shading)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.cell(ri + 1, ci).text = str(val)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)


def create_report():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ═══════ 封面 ═══════
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Flask 用户管理系统')
    run.bold = True; run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('XXE 漏洞检测与修复报告')
    run.bold = True; run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x66, 0x7e, 0xea)

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run('━' * 40)
    run.font.color.rgb = RGBColor(0x66, 0x7e, 0xea)
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        '项目地址：github.com/nobody848/flask-user-management\n'
        '报告日期：2026-07-16\n'
        '测试范围：/xml-import 路由 XXE 漏洞检测与修复\n'
        '修复提交：当前分支最新'
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    # ══════════════════════════════════════
    # 一、XXE 漏洞原理
    # ══════════════════════════════════════
    doc.add_heading('一、XXE 漏洞原理', level=1)

    doc.add_paragraph(
        'XXE（XML External Entity，XML 外部实体注入）是一种 XML 解析安全漏洞。'
        '当 XML 解析器允许解析外部实体时，攻击者可以通过构造恶意 XML 内容，'
        '导致服务器读取本地文件、发起内网请求（SSRF）或执行拒绝服务攻击。'
    )

    doc.add_paragraph('XXE 攻击的常见危害：')
    buls = [
        '任意文件读取：通过 file:// 协议读取服务器上的敏感文件（/etc/passwd、配置文件等）',
        'SSRF 内网探测：通过 http:// 协议访问内网服务',
        '拒绝服务攻击：通过实体嵌套（Billion Laughs 攻击）耗尽服务器资源',
        '文件解析绕过：利用 UTF-7 编码等方式绕过输入过滤',
    ]
    for b in buls:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph('XXE 攻击的基本 XML 结构：')
    add_code(doc, '''<?xml version="1.0"?>
<!DOCTYPE foo [
  <!ENTITY xxe SYSTEM "file:///etc/passwd">
]>
<root>&xxe;</root>
<!-- 解析后 &xxe; 被替换为 /etc/passwd 的文件内容 -->''')

    doc.add_page_break()

    # ══════════════════════════════════════
    # 二、漏洞发现
    # ══════════════════════════════════════
    doc.add_heading('二、漏洞发现', level=1)

    doc.add_heading('2.1 功能概述', level=2)
    doc.add_paragraph(
        'XML 数据导入功能（/xml-import 路由）允许登录用户提交 XML 数据，'
        '解析其中的 user 节点并以 JSON 格式返回 name 和 email 字段。'
    )

    doc.add_heading('2.2 漏洞代码定位（修复前）', level=2)
    add_code(doc, '''# ❌ 修复前（app.py 第695-712行）：
# 直接提取 SYSTEM 声明的文件路径并 open() 读取
entity_pattern = re.compile(r'<!ENTITY\\s+\\w+\\s+SYSTEM\\s+"([^"]+)"')
file_paths = entity_pattern.findall(xml_data)

for fpath in file_paths:
    with open(fpath, "r", encoding="utf-8") as f:
        file_content = f.read()     # ⚠️ 读取任意文件
    ename = ...
    replaced_xml = replaced_xml.replace(f"&{ename};", file_content)''')

    doc.add_heading('2.3 漏洞分析', level=2)

    doc.add_heading('XXE-01：DTD 外部实体未禁用（高危）', level=3)
    doc.add_paragraph(
        '代码显式查找 `<!ENTITY ... SYSTEM "...">` 声明，提取文件路径后直接调用 '
        'open() 读取文件内容。攻击者只需构造含有 SYSTEM 实体声明的 XML，'
        '即可读取服务器上的任意文件。'
    )
    add_code(doc, '''# 攻击者提交的 payload：
<!DOCTYPE foo [
  <!ENTITY xxe SYSTEM "/etc/passwd">
]>
<users>
    <user>
        <name>&xxe;</name>      <!-- 被替换为 /etc/passwd 内容 -->
        <email>test@test.com</email>
    </user>
</users>

# 返回结果中 name 字段包含 /etc/passwd 的完整内容''')

    doc.add_heading('XXE-02：无文件路径白名单（高危）', level=3)
    doc.add_paragraph(
        '代码不对文件路径做任何校验，攻击者可以读取系统中的任意文件。'
    )

    doc.add_heading('XXE-03：无外部实体数量限制（低危）', level=3)
    doc.add_paragraph(
        '代码未限制 ENTITY 声明数量，攻击者可构造递归实体导致 DoS。'
    )

    doc.add_page_break()

    # ══════════════════════════════════════
    # 三、攻击路径推演
    # ══════════════════════════════════════
    doc.add_heading('三、攻击路径推演', level=1)

    doc.add_heading('3.1 任意文件读取攻击', level=2)
    add_table(doc,
        ['步骤', '操作', '说明'],
        [
            ['1', '攻击者登录系统', '获取合法 session'],
            ['2', '构造 XXE payload', '<!ENTITY xxe SYSTEM "/etc/passwd">'],
            ['3', 'POST /xml-import', '提交包含 XXE 的 XML'],
            ['4', '服务器读取文件', 'open("/etc/passwd") 读取本地文件'],
            ['5', '返回文件内容', 'name 字段包含 /etc/passwd 内容'],
        ],
        col_widths=[1, 4, 8.5]
    )

    doc.add_heading('3.2 可读取的敏感文件', level=2)
    add_table(doc,
        ['文件路径', '内容说明'],
        [
            ['/etc/passwd', '系统用户列表'],
            ['/etc/shadow', '用户密码哈希（需 root 权限）'],
            ['/etc/hosts', '主机名解析配置'],
            ['/app.py', '项目源代码（当前目录）'],
            ['/proc/self/environ', '环境变量（可能含密钥）'],
            ['/data/users.db', 'SQLite 数据库文件'],
        ],
        col_widths=[5.5, 8]
    )

    doc.add_page_break()

    # ══════════════════════════════════════
    # 四、修复方案
    # ══════════════════════════════════════
    doc.add_heading('四、修复方案', level=1)

    doc.add_heading('4.1 修复思路', level=2)
    doc.add_paragraph(
        '最彻底的修复方式是：在 XML 解析前移除 DOCTYPE 声明及其内部所有子集（DTD），'
        '从源头阻断外部实体声明。这样既保留了正常的 XML 数据解析功能，'
        '又彻底杜绝了 XXE 攻击向量。'
    )

    doc.add_heading('4.2 修复代码', level=2)
    add_code(doc, '''# ✅ 修复后（app.py 第695-709行）：
# XXE 防护：移除 DOCTYPE 声明及其内部子集（含 ENTITY 定义）
safe_xml = re.sub(
    r'<!DOCTYPE\\s+\\w+\\s*\\[.*?\\]\\s*>',
    '',
    xml_data,
    flags=re.DOTALL
)

# 使用正则安全提取 user 节点（不加载 XML 解析器）
users = []
for m in re.finditer(
    r'<user>\\s*<name>(.*?)</name>\\s*<email>(.*?)</email>\\s*</user>',
    safe_xml, re.DOTALL
):
    users.append({
        "name": m.group(1).strip(),
        "email": m.group(2).strip()
    })''')

    doc.add_heading('4.3 防护原理', level=2)
    add_table(doc,
        ['攻击方式', '修复前', '修复后'],
        [
            ['<!ENTITY xxe SYSTEM "file:///etc/passwd">', '❌ 读取文件', '✅ DOCTYPE 被移除'],
            ['<!ENTITY xxe SYSTEM "http://169.254.169.254/">', '❌ SSRF', '✅ DOCTYPE 被移除'],
            ['<!ENTITY lol "lol"><!ENTITY lol2 "&lol;&lol;">...', '❌ DoS', '✅ DOCTYPE 被移除'],
            ['<!ENTITY xxe SYSTEM "\\\\192.168.1.1\\share">', '❌ SMB 攻击', '✅ DOCTYPE 被移除'],
            ['正常 XML（无 DTD）', '✅ 正常解析', '✅ 正常解析'],
        ],
        col_widths=[6, 3, 3.5]
    )

    doc.add_page_break()

    # ══════════════════════════════════════
    # 五、修复验证
    # ══════════════════════════════════════
    doc.add_heading('五、修复验证', level=1)

    add_table(doc,
        ['编号', '测试用例', '输入', '预期', '结果'],
        [
            ['T-01', '正常 XML 解析', '普通 <user> 节点', '返回 JSON', '✅'],
            ['T-02', 'XXE 文件读取', 'SYSTEM "/etc/passwd"', '被拦截，不返回文件内容', '✅'],
            ['T-03', '安全 ENTITY', '内部实体无 SYSTEM', '正常解析', '✅'],
            ['T-04', '空 XML 提交', '空字符串', '不报错', '✅'],
            ['T-05', '未登录访问', 'GET /xml-import', '302 跳转', '✅'],
        ],
        col_widths=[1.5, 3.5, 4, 3.5, 1.5]
    )

    doc.add_paragraph()

    # ══════════════════════════════════════
    # 六、漏洞汇总
    # ══════════════════════════════════════
    doc.add_heading('六、漏洞汇总', level=1)

    add_table(doc,
        ['编号', '漏洞名称', '严重', '发现位置', '状态'],
        [
            ['XXE-01', 'DTD 外部实体未禁用（任意文件读取）', '🔴 高危', '/xml-import 第695-704行', '✅ 已修复'],
            ['XXE-02', '文件路径无白名单校验', '🔴 高危', '/xml-import 第697行', '✅ 已修复'],
            ['XXE-03', '无外部实体数量限制（DoS 风险）', '🟡 低危', '/xml-import 第696行', '✅ 已修复'],
        ],
        col_widths=[2, 6, 2, 3.5, 2.5]
    )

    doc.add_paragraph()

    # ══════════════════════════════════════
    # 七、Git 提交记录
    # ══════════════════════════════════════
    doc.add_heading('七、Git 提交记录', level=1)

    add_table(doc,
        ['提交哈希', '提交信息', '涉及文件'],
        [
            ['a375226', 'feat: 新增XML数据导入功能', 'app.py, xml_import.html, base.html'],
            ['当前', 'fix: 修复XXE漏洞（XML数据导入）', 'app.py'],
        ],
        col_widths=[3, 6, 5]
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 40)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('报告生成日期：2026-07-16')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── 保存 ──
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(script_dir, 'static', 'XXE漏洞检测与修复报告.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f'报告已生成：{output_path}')
    print(f'文件大小：{os.path.getsize(output_path) / 1024:.1f} KB')
    return output_path


if __name__ == '__main__':
    create_report()
